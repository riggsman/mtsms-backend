"""
Helper module for generating PDF and Word documents from notes
"""
import os
from typing import Optional
from datetime import datetime

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from bs4 import BeautifulSoup
import re

def clean_html(html_content: str) -> str:
    """Remove HTML tags and clean content"""
    if not html_content:
        return ""
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text()

def generate_pdf(
    title: str,
    course_name: str,
    course_code: Optional[str],
    content: str,
    lecturer_name: Optional[str] = None,
    output_path: str = None
) -> Optional[str]:
    """
    Generate a PDF file from note content
    
    Returns:
        Path to generated PDF file or None if generation fails
    """
    if not REPORTLAB_AVAILABLE:
        print("ReportLab not available. Install it with: pip install reportlab")
        return None
    
    try:
        # Create output directory if it doesn't exist
        if output_path:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
        else:
            # Default output directory
            output_dir = os.path.join(os.getcwd(), 'uploads', 'notes', 'pdf')
            os.makedirs(output_dir, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_title = re.sub(r'[^\w\s-]', '', title)[:50]
            filename = f"{safe_title}_{timestamp}.pdf"
            output_path = os.path.join(output_dir, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#000000',
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#333333',
            spaceAfter=8,
            alignment=TA_LEFT,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            textColor='#000000',
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add course information
        course_info = f"<b>Course:</b> {course_name}"
        if course_code:
            course_info += f" ({course_code})"
        story.append(Paragraph(course_info, normal_style))
        story.append(Spacer(1, 0.1*inch))
        
        if lecturer_name:
            story.append(Paragraph(f"<b>Lecturer:</b> {lecturer_name}", normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%B %d, %Y')}", normal_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Add content
        # Clean HTML and convert to paragraphs
        cleaned_content = clean_html(content)
        paragraphs = cleaned_content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                # Replace newlines within paragraphs with <br/>
                para = para.replace('\n', '<br/>')
                story.append(Paragraph(para, normal_style))
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        
        return output_path
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return None

def generate_word(
    title: str,
    course_name: str,
    course_code: Optional[str],
    content: str,
    lecturer_name: Optional[str] = None,
    output_path: str = None
) -> Optional[str]:
    """
    Generate a Word document from note content
    
    Returns:
        Path to generated Word file or None if generation fails
    """
    if not DOCX_AVAILABLE:
        print("python-docx not available. Install it with: pip install python-docx")
        return None
    
    try:
        # Create output directory if it doesn't exist
        if output_path:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
        else:
            # Default output directory
            output_dir = os.path.join(os.getcwd(), 'uploads', 'notes', 'word')
            os.makedirs(output_dir, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_title = re.sub(r'[^\w\s-]', '', title)[:50]
            filename = f"{safe_title}_{timestamp}.docx"
            output_path = os.path.join(output_dir, filename)
        
        # Create Word document
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add course information
        doc.add_paragraph()
        course_info = f"Course: {course_name}"
        if course_code:
            course_info += f" ({course_code})"
        doc.add_paragraph(course_info, style='Intense Quote')
        
        if lecturer_name:
            doc.add_paragraph(f"Lecturer: {lecturer_name}", style='Intense Quote')
        
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", style='Intense Quote')
        doc.add_paragraph()
        
        # Add content
        # Clean HTML and convert to paragraphs
        cleaned_content = clean_html(content)
        paragraphs = cleaned_content.split('\n\n')
        
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save document
        doc.save(output_path)
        
        return output_path
    except Exception as e:
        print(f"Error generating Word document: {e}")
        return None
